"""
fill_valyuta.py — заповнює оригінальний шаблон через python-docx,
конвертує у PDF через LibreOffice напряму.
"""
import copy, io, os, re, shutil, subprocess, tempfile
from datetime import date
from docx import Document
from docx.oxml.ns import qn

MONTHS_UK = {
    1:"січня",2:"лютого",3:"березня",4:"квітня",
    5:"травня",6:"червня",7:"липня",8:"серпня",
    9:"вересня",10:"жовтня",11:"листопада",12:"грудня",
}
CURRENCY_CODES    = {"USD":"840","EUR":"978"}
CURRENCY_NAMES_UK = {"USD":"Долар США","EUR":"Євро"}


def get_date_parts(date_str):
    if date_str and date_str.strip():
        p = date_str.strip().split(".")
        d,m,y = int(p[0]),int(p[1]),int(p[2])
    else:
        t = date.today(); d,m,y = t.day,t.month,t.year
    return str(d), MONTHS_UK[m], str(y)


def set_cell_text(cell, new_text):
    """Замінює текст у клітинці зберігаючи форматування першого run."""
    for para in cell.paragraphs:
        if not para.runs:
            continue
        # Зберігаємо форматування першого run
        first_run = para.runs[0]
        rpr = copy.deepcopy(first_run._r.find(qn('w:rPr')))
        # Очищаємо всі runs
        for run in para.runs:
            run._r.getparent().remove(run._r)
        # Додаємо один новий run
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
        t = OxmlElement('w:t')
        t.text = new_text
        if new_text.startswith(' ') or new_text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        para._p.append(r)
        return  # тільки перший параграф


def find_cell_with_text(table, text):
    """Знаходить першу клітинку з точним текстом."""
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == text:
                return cell
    return None


def replace_in_cell(table, old_text, new_text):
    """Замінює текст у всіх клітинках таблиці де він співпадає."""
    replaced = 0
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == old_text:
                set_cell_text(cell, new_text)
                replaced += 1
    return replaced


def replace_iban_in_table(table, old_iban, new_iban):
    """Замінює цифри IBAN в таблиці (кожна цифра — окрема клітинка)."""
    old_chars = list(old_iban.replace(" ",""))
    new_chars = list(new_iban.replace(" ","").ljust(len(old_chars)))[:len(old_chars)]
    # Збираємо всі клітинки таблиці в порядку
    all_cells = []
    for row in table.rows:
        for cell in row.cells:
            all_cells.append(cell)
    # Шукаємо послідовність
    texts = [c.text.strip() for c in all_cells]
    for i in range(len(texts) - len(old_chars)):
        if texts[i:i+len(old_chars)] == old_chars:
            for j, ch in enumerate(new_chars):
                set_cell_text(all_cells[i+j], ch)
            return True
    return False


def replace_amount_in_table(table, old_digits, new_amount_str):
    """Замінює цифри суми в таблиці."""
    try:
        val = float(new_amount_str)
        int_p = str(int(val))
        dec_p = f"{val:.2f}".split(".")[1]
        new_chars = list(int_p) + [","] + list(dec_p)
    except Exception:
        return False
    old_chars = old_digits
    # Вирівнюємо довжину
    while len(new_chars) < len(old_chars):
        new_chars.insert(0, " ")
    new_chars = new_chars[-len(old_chars):]

    all_cells = []
    for row in table.rows:
        for cell in row.cells:
            all_cells.append(cell)
    texts = [c.text.strip() for c in all_cells]
    for i in range(len(texts) - len(old_chars)):
        if texts[i:i+len(old_chars)] == old_chars:
            for j, ch in enumerate(new_chars):
                set_cell_text(all_cells[i+j], ch if ch.strip() else " ")
            return True
    return False


def replace_run_text(doc, old_text, new_text):
    """Замінює текст у всіх параграфах документа (поза таблицями)."""
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)


def replace_date_in_header(table, day, month_uk, year):
    """Замінює дату у рядку заголовка."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full = para.text
                if 'ВІД/DATED' in full and 'березня' in full:
                    # Замінюємо у runs
                    for run in para.runs:
                        run.text = run.text.replace('23', day, 1) if '23' in run.text else run.text
                        run.text = run.text.replace('березня', month_uk) if 'березня' in run.text else run.text
                        run.text = run.text.replace('2026', year) if '2026' in run.text else run.text
                    return True
    return False


def build_zayava_pdf(params):
    """Заповнює шаблон і повертає PDF bytes."""
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                  "template_valyuta.docx")
    tmpdir = tempfile.mkdtemp()
    out_docx = os.path.join(tmpdir, "zayava.docx")
    out_pdf  = os.path.join(tmpdir, "zayava.pdf")

    pib         = params["pib"]
    pib_parts   = pib.split()
    ipn         = params["ipn"]
    address     = params["address"]
    purpose_uk  = params["purpose_uk"]
    amount      = params["amount"]
    currency    = params["currency"]
    cur_code    = CURRENCY_CODES.get(currency,"840")
    cur_name_uk = CURRENCY_NAMES_UK.get(currency,"Долар США")
    iban_debit  = params["iban_debit"].replace(" ","")
    iban_credit = params["iban_credit"].replace(" ","")
    iban_comm   = params["iban_commission"].replace(" ","")
    day, month_uk, year = get_date_parts(params.get("date",""))

    doc = Document(template_path)
    tables = doc.tables

    # ── T0: Заголовок, дата, клієнт ─────────────────────────────────────────
    t0 = tables[0]

    # Дата
    replace_date_in_header(t0, day, month_uk, year)

    # ПІБ клієнта
    for row in t0.rows:
        for cell in row.cells:
            txt = cell.text.strip()
            if txt == "ФОП Каллаш Леонід Юрійович":
                set_cell_text(cell, f"ФОП {pib}")
            elif "Костя Гордієнка" in txt or txt == "Україна, м. Київ, пров. Костя Гордієнка, буд. 10, кв. 7":
                set_cell_text(cell, address)
            elif txt == "3073810850":
                set_cell_text(cell, ipn)

    # ── T1: Мета купівлі ─────────────────────────────────────────────────────
    t1 = tables[1]
    for row in t1.rows:
        for cell in row.cells:
            txt = cell.text.strip()
            if txt not in ("Мета купівлі/ Purpose\nof purchase",
                           "Мета купівлі/ Purpose", "of purchase") \
               and "Мета купівлі" not in txt and len(txt) > 10:
                set_cell_text(cell, purpose_uk)

    # ── T2: Сума, валюта, IBAN 1 ─────────────────────────────────────────────
    t2 = tables[2]

    # Сума — старі цифри ["5","8","0","0",",","0","0"]
    replace_amount_in_table(t2, ["5","8","0","0",",","0","0"], amount)

    # Код валюти — старі ["8","4","0"]
    old_cur = list(CURRENCY_CODES.get("USD","840"))
    new_cur = list(cur_code.ljust(3))
    all_cells = []
    for row in t2.rows:
        for cell in row.cells:
            all_cells.append(cell)
    texts = [c.text.strip() for c in all_cells]
    for i in range(len(texts)-3):
        if texts[i:i+3] == old_cur:
            for j,ch in enumerate(new_cur):
                set_cell_text(all_cells[i+j], ch)
            break

    # Назва валюти
    for row in t2.rows:
        for cell in row.cells:
            if cell.text.strip() == "Долар США":
                set_cell_text(cell, cur_name_uk)
                break

    # IBAN 1 (в T2)
    replace_iban_in_table(t2,
        "UA293220010000026006310115156",
        iban_debit)

    # ── T3: IBAN 2 ───────────────────────────────────────────────────────────
    replace_iban_in_table(tables[3],
        "UA353220010000026000370076190",
        iban_credit)

    # ── T4: IBAN 3 ───────────────────────────────────────────────────────────
    replace_iban_in_table(tables[4],
        "UA293220010000026006310115156",
        iban_comm)

    # ── T5: Підпис (ПІБ внизу) ───────────────────────────────────────────────
    t5 = tables[5]
    for row in t5.rows:
        for cell in row.cells:
            if "Каллаш Леонід Юрійович" in cell.text:
                # Зберігаємо "М.П./Seal\n" + новий ПІБ
                for para in cell.paragraphs:
                    if "Каллаш" in para.text:
                        for run in para.runs:
                            if "Каллаш" in run.text or "Леонід" in run.text or "Юрійович" in run.text:
                                run.text = run.text.replace("Каллаш", pib_parts[0] if len(pib_parts)>0 else "")
                                run.text = run.text.replace("Леонід", pib_parts[1] if len(pib_parts)>1 else "")
                                run.text = run.text.replace("Юрійович", pib_parts[2] if len(pib_parts)>2 else "")

    doc.save(out_docx)

    # ── Конвертація docx → PDF через LibreOffice ─────────────────────────────
    def _find_lo():
        # 1. Пряма перевірка відомих шляхів
        for path in ["/usr/bin/libreoffice", "/usr/bin/soffice",
                     "/usr/lib/libreoffice/program/soffice",
                     "/opt/libreoffice7.6/program/soffice",
                     "/opt/libreoffice/program/soffice"]:
            if os.path.isfile(path):
                return path
        # 2. Через shutil.which (залежить від PATH)
        for name in ["libreoffice", "soffice"]:
            found = shutil.which(name)
            if found:
                return found
        # 3. Через find
        try:
            r = subprocess.run(["find", "/usr", "/opt", "-name", "soffice",
                                 "-type", "f"], capture_output=True, text=True, timeout=10)
            for line in r.stdout.strip().splitlines():
                if line.strip():
                    return line.strip()
        except Exception:
            pass
        return None

    lo_bin = _find_lo()
    if not lo_bin:
        raise RuntimeError("LibreOffice не знайдено. Встановіть: apt-get install libreoffice")

    env = os.environ.copy()
    env["HOME"] = tmpdir

    subprocess.run(
        [lo_bin, "--headless", "--norestore",
         "--convert-to", "pdf", "--outdir", tmpdir, out_docx],
        check=True, capture_output=True, timeout=90, env=env
    )

    # LibreOffice зберігає як zayava.pdf
    if not os.path.exists(out_pdf):
        # Шукаємо будь-який pdf у tmpdir
        for f in os.listdir(tmpdir):
            if f.endswith(".pdf"):
                out_pdf = os.path.join(tmpdir, f)
                break

    with open(out_pdf, "rb") as f:
        data = f.read()

    shutil.rmtree(tmpdir, ignore_errors=True)
    return data


# Заглушки для сумісності
def fill_template(params, template_path, output_docx): pass
def docx_to_pdf(docx_path, output_pdf): pass


if __name__ == "__main__":
    params = {
        "pib": "Мацола Євгеній Володимирович",
        "ipn": "3522908011",
        "address": "Україна, обл. Закарпатська, р-н. Тячівський, с. Грушово, вул. Центральна, буд. 53-А",
        "date": "",
        "purpose_uk": "Передоплата за товар згідно з договором № 01/26 від 16.03.2026, рахунком № PI2026031201 від 17.03.2026",
        "amount": "5800.00", "currency": "USD",
        "iban_debit":      "UA293220010000026006310115156",
        "iban_credit":     "UA353220010000026000370076190",
        "iban_commission": "UA293220010000026006310115156",
    }
    print("Генеруємо...")
    pdf = build_zayava_pdf(params)
    with open("/tmp/test_final.pdf","wb") as f:
        f.write(pdf)
    print(f"✅ PDF: {len(pdf)} bytes → /tmp/test_final.pdf")
